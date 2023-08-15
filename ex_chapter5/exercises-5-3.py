#!/usr/bin/python3
################################################################################
# 5.3 Parse FastQC results, create a summary report as .docx including image
################################################################################
# Import required modules
import os
import re
from docx import Document
from docx.shared import Cm

################################################################################
# V1
# Setting working directory
location = os.getcwd()
ex53dir = location + "/HTA/fastqc/"
# Opening new document
document = Document()
# Start scanning working directory for folders to read
print("Scanning files and folders in {}".format(ex53dir))
file_list = os.listdir(ex53dir)
for folder in file_list:
    path_folder = ex53dir + folder
    if(os.path.isdir(path_folder)):
        print("Reading directory {}".format(path_folder))
        # Read fastqc_data.txt
        # Read summary.txt and write in .docx report
        # Add image
        # Adding a page break between samples
# Save document
document.save('fastqc_summary_V1.docx')



################################################################################
# V2

# Functions
def get_summary(file):
    fileObj = open(file, "r")
    d = {}
    for line in fileObj.readlines():
        fields = line.split("\t")
        d[fields[1]] = fields[0]
    return d
    
# Setting working directory
location = os.getcwd()
ex53dir = location + "/HTA/fastqc/"

# Opening new document
document = Document()

# Start scanning working directory for folders to read
print("Scanning files and folders in {}".format(ex53dir))
file_list = os.listdir(ex53dir)
for folder in file_list:
    path_folder = ex53dir + folder
    if(os.path.isdir(path_folder)):
        print("Reading directory {}".format(path_folder))
        sample = folder.split('_')
        sample_heading = "Sample " + sample[0]
        document.add_heading(sample_heading, level=1)
        document.add_heading("Summary", level=3)
        # Read fastqc_data.txt
        data_file = path_folder + "/fastqc_data.txt"
        fileObj = open(data_file, "r")
        for line in fileObj.readlines():
            total_seq = re.search("^Total", line)
            if total_seq:
                p = document.add_paragraph(line.rstrip())
        # Read summary.txt and write in .docx report
        summary_file = path_folder + "/summary.txt"
        summary_dict = get_summary(summary_file)
        for k, v in summary_dict.items():
            if(k=="Basic Statistics" or k=="Per base sequence quality"):
                summary_result = k + ": " + v
                p = document.add_paragraph(summary_result)
        # Add image
        quality_img = path_folder + "/Images/per_base_quality.png"
        document.add_picture(quality_img, width=Cm(12))
        # Adding a page break between samples
        document.add_page_break()

# Save document
document.save('fastqc_summary.docx')
