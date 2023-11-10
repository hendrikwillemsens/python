#!/usr/bin/python3
import csv
import MySQLdb as my
from docx import Document
document = Document()
db = my.connect(host="genome-mysql.soe.ucsc.edu",

   user="genomep",
   passwd="password",
   db="hg19")
c = db.cursor()

print("Reading regions from file:")
print("==========================")

with open('coordinates.csv') as csv_file:
    csv_reader = csv.reader(csv_file, delimiter=',')
    line_count = 0
    for row in csv_reader:
        chromosoom = row[0].split(":")
        positie = chromosoom[1].split("-")
        print("chrom = {}       start = {}      end = {}".format(chromosoom[0], positie[0], positie[1]))
        print("getting")
        rows = c.execute("""SELECT name, chrom, chromStart, chromEnd, score, strand FROM targetScanS where chrom = '{}' AND chromStart > '{}' AND chromEnd < '{}'""".format(chromosoom[0],positie[0], positie[1]))
        print("Total rows ensGene: {}".format(rows))
        result = c.fetchall()
        document.add_heading('Ts miRNA sites for {}'.format(row[0]), level=1)
        table = document.add_table(rows=rows + 1, cols=5)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = 'namegene'
        hdr_cells[1].text = 'name miRNA'
        hdr_cells[2].text = 'positie'
        hdr_cells[3].text = 'score'
        hdr_cells[4].text = 'strand'
        
        




        document.add_page_break()
        
        





        

db.close()

document.save('oef1.docx')

        

    

csv_file.close()