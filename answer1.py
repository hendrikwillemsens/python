#!/usr/bin/python3
import re
import csv
import MySQLdb as my
from docx import Document

document = Document()
# Reading csv file line by line
print("Reading regions from file")   
print("=" *20)  
with open('coordinates.csv') as csv_file:
    csv_reader = csv.reader(csv_file)
    for row in csv_reader:
        print(', '.join(row))
        rij = row[0].split(":")
        '''print(rij)'''''
        rijen = rij[1].split("-")
        '''print(rijen)'''
        print("chr:{}\tstart:{}\t\tend:{}".format(rij[0],rijen[0],rijen[1]))
        print("getting TS from mirRNA sites using query:")
        print("""SELECT * from targetScanS WHERE chrom = '{}' AND chromstart >'{}'AND chromEnd < '{}'.""".format(rij[0],rijen[0],rijen[1]))
        
        
        db = my.connect(host="genome-mysql.soe.ucsc.edu",
        user="genomep",
        passwd="password",
        db="hg19")
        c = db.cursor()
        no_rows = c.execute("""SELECT * from targetScanS WHERE chrom = '{}' AND chromStart > '{}' AND chromEnd < '{}'""".format(rij[0],rijen[0],rijen[1]))      
        print("Result:{}".format(no_rows))
        print("\n")
        
        document.add_heading('TS miRNA sites for {}'.format(', '.join(row)), 0)
        tabel = no_rows + 1
        table = document.add_table(rows=tabel , cols=5)
        table.style = 'Light Shading Accent 1'
        # populate header row --------
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Name gene'
        heading_cells[1].text = 'Name miRNA'
        heading_cells[2].text = 'Postion'
        heading_cells[3].text = 'Score'
        heading_cells[4].text = 'Strand'
        document.add_page_break()

       
        

db.close()
csv_file.close()
document.save('answer1.docx')
        
 


